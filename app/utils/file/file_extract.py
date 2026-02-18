import fitz  # PyMuPDF
from docx import Document

def extract_text_from_pdf(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

# def extract_text_from_docx(docx_path: str) -> str:
#     doc = Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    sections = []

    # 1. 提取段落和标题
    for para in doc.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()

        if not text:
            continue

        # 根据样式分类：标题 or 段落
        if "heading" in style:
            level = ''.join(filter(str.isdigit, style)) or "1"
            sections.append(f"[标题{level}] {text}")
        else:
            sections.append(f"[段落] {text}")

    # 2. 提取表格内容
    for table_idx, table in enumerate(doc.tables, 1):
        sections.append(f"[表格{table_idx}]")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
            if row_text:
                sections.append(row_text)
        sections.append("")  # 表格之间空行

    return "\n".join(sections)